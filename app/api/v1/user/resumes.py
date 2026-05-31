import os
import re
import shutil
import uuid
from typing import List, Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session

# --- 1. 数据库与模型导入 ---
from app.models.database import get_db
from app.models.resume import Resume
from app.models.user import User
from app.models import base

# --- 2. Schema 导入 ---
from app.schemas.resume_schema import ResumeCreate, ResumeUpdate, ResumeRead, ResumeBase

# --- 3. 核心组件导入 ---
from app.core import deps
from app.core.llm import get_llm  # 导入获取 LLM 的工厂函数
from app.core.vector_store import add_documents_to_vectorstore
from app.core.rag_engine import process_file_pipeline

# --- 4. LangChain 组件导入 (仅用于解析) ---
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate

import io
import pdfplumber
from docx import Document

router = APIRouter(prefix="/api/v1/resume", tags=["个人简历管理"])

# --- 配置常量 ---
TEMP_UPLOAD_DIR = "temp_uploads"
os.makedirs(TEMP_UPLOAD_DIR, exist_ok=True)

# ==========================================
#           辅助函数区域
# ==========================================


def concatenate_resume_text(data: ResumeCreate) -> str:
    """
    【向量库策略】将核心字段拼接成一段长文本
    """
    parts = []
    if data.summary:
        parts.append(f"【个人评价】{data.summary}")
    if data.work_experience:
        parts.append(f"【工作经历】{data.work_experience}")
    if data.project_experience:
        parts.append(f"【项目经历】{data.project_experience}")
    return "\n".join(parts)


def fallback_extract(text: str) -> dict:
    """
    【解析降级策略】当 LLM 不可用时使用的正则提取
    """
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    phone_match = re.search(r"1[3-9]\d{9}", text)
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    name = lines[0] if lines else ""

    return {
        "name": name,
        "phone": phone_match.group(0) if phone_match else "",
        "email": email_match.group(0) if email_match else "",
        "title": "未命名简历",
        "summary": text[:200] + "...",
        "work_experience": text,
        "project_experience": "",
        "education": "",
    }


async def llm_extract_resume(text: str) -> dict:
    """
    【解析核心】使用 LLM 从非结构化文本中提取简历信息
    """
    try:
        # 1. 获取 LLM 实例 (修复：不再使用全局变量 llm，而是调用函数)
        model = get_llm()

        # 2. 定义提示词
        prompt_template = """
你是一个资深技术招聘专家，擅长从杂乱的简历文本中提取关键信息。
请仔细阅读下方的简历文本，提取信息并转换为严格的 JSON 格式。

**提取规则：**
1. **技能清洗**：在提取 `skills` 时，请去除“精通”、“熟悉”、“掌握”等修饰词，只保留技术名词（例如：将“精通Java”提取为“Java”）。
2. **时间标准化**：
   - `start_date` 和 `end_date` 必须统一为 `YYYY-MM` 格式。
   - 如果遇到“至今”、“现在”，结束日期请填写 `2026-05`（假设当前时间）。
   - 如果只有年份（如2020年），月份默认补全为 `01`。
3. **完整性**：`project_experience` 中的 `description` 和 `responsibilities` 尽量保留原文的关键细节，不要过度概括。
4. **格式**：输出必须是纯净的 JSON 字符串，不要包含 Markdown 格式（如 ```json），也不要包含任何解释性文字。

**JSON 结构定义：**
{{
    "name": "姓名",
    "phone": "手机号",
    "email": "邮箱",
    "education": [
        {{
            "school": "学校",
            "major": "专业",
            "degree": "学历",
            "start_date": "YYYY-MM",
            "end_date": "YYYY-MM"
        }}
    ],
    "work_experience": [
        {{
            "company": "公司",
            "position": "职位",
            "start_date": "YYYY-MM",
            "end_date": "YYYY-MM",
            "description": "工作内容（简练）"
        }}
    ],
    "project_experience": [
        {{
            "name": "项目名",
            "role": "角色",
            "start_date": "YYYY-MM",
            "end_date": "YYYY-MM",
            "description": "项目详情",
            "technologies": ["用到的技术栈"]
        }}
    ],
    "skills": {{
        "languages": ["语言"],
        "frameworks": ["框架"],
        "databases": ["数据库"],
        "tools": ["工具"]
    }}
}}

**简历文本如下：**
{text}
"""
        prompt = PromptTemplate(template=prompt_template, input_variables=["text"])
        parser = JsonOutputParser(pydantic_object=ResumeBase)

        # 3. 构建链并执行
        chain = prompt | model | parser
        result = chain.invoke({"text": text[:3000]})  # 限制长度
        return result

    except Exception as e:
        print(f"LLM 解析失败: {e}")
        return fallback_extract(text)


# ==========================================
#           API 路由区域
# ==========================================


# --- 1. 获取列表 ---
@router.get("/list", response_model=List[ResumeRead])
def get_resumes(
    db: Session = Depends(get_db), current_user: User = Depends(deps.get_current_user)
):
    """获取当前用户的简历列表"""
    resumes = db.query(Resume).filter(Resume.user_id == current_user.id).all()
    return resumes


import traceback  # 1. 新增导入：用于打印详细的错误堆栈
import logging  # 1. 新增导入：用于记录日志

# 假设你已经在文件头部配置了 logger，如果没有，可以使用 print 或者下面这行
logger = logging.getLogger(__name__)

from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.exceptions import OutputParserException


# --- 核心辅助函数：从文件名中提取关键信息 ---
def extract_info_from_filename(filename: str) -> dict:
    """
    从文件名中提取姓名、职位、薪资等信息
    支持常见的命名格式，如：张三_前端开发_15k.pdf、李四-Java工程师-3年经验.docx
    """
    # 去掉文件后缀
    name_without_ext = os.path.splitext(filename)[0]
    # 替换常见的分隔符（下划线、横杠、空格）为统一的空格
    clean_name = re.sub(r"[_\-\s]+", " ", name_without_ext).strip()

    extracted_info = {
        "filename_hint": clean_name,
        "possible_name": "",
        "possible_title": "",
        "possible_salary": "",
    }

    # 1. 提取薪资（匹配如：15k, 20K, 15-25k, 20k-30k 等）
    salary_match = re.search(
        r"(\d+\.?\d*)\s*[-~至到]?\s*(\d*)\s*k", clean_name, re.IGNORECASE
    )
    if salary_match:
        min_s = salary_match.group(1)
        max_s = salary_match.group(2)
        if max_s:
            extracted_info["possible_salary"] = f"{min_s}k-{max_s}k"
        else:
            extracted_info["possible_salary"] = f"{min_s}k"

    # 2. 提取职位（这里做一个简单的启发式匹配，实际可以根据你的业务扩展关键词库）
    # 匹配文件名中常见的职位关键词
    job_keywords = [
        "前端",
        "后端",
        "Java",
        "Python",
        "算法",
        "产品",
        "运营",
        "测试",
        "UI",
        "设计",
        "工程师",
        "开发",
        "经理",
        "总监",
        "实习",
    ]
    found_jobs = [kw for kw in job_keywords if kw.lower() in clean_name.lower()]
    if found_jobs:
        # 简单取第一个匹配到的，或者你可以写更复杂的逻辑提取完整职位短语
        extracted_info["possible_title"] = found_jobs[0]

    # 3. 提取姓名（假设文件名开头2-4个字符是中文姓名）
    name_match = re.match(r"^([\u4e00-\u9fa5]{2,4})", clean_name)
    if name_match:
        extracted_info["possible_name"] = name_match.group(1)

    return extracted_info


# --- 辅助函数：提取文档纯文本（复用之前的逻辑） ---
def extract_text_from_file(file: UploadFile) -> str:
    contents = file.file.read()
    full_text = ""
    try:
        if file.filename.lower().endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
        elif file.filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(contents))
            full_text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if row_text:
                        full_text_parts.append(" | ".join(row_text))
            full_text = "\n".join(full_text_parts)
        elif file.filename.lower().endswith(".txt"):
            full_text = contents.decode("utf-8")
        if not full_text.strip():
            raise ValueError("无法提取文档文本")
        return full_text
    finally:
        file.file.seek(0)


# --- 简历智能解析接口（升级版） ---
@router.post("/parse")
async def parse_resume_document(
    file: UploadFile = File(...),
    # current_user: UserDB = Depends(get_current_active_admin)
):
    if not file.filename.lower().endswith((".pdf", ".docx", ".txt")):
        raise HTTPException(status_code=400, detail="仅支持 PDF, DOCX 或 TXT 文件")

    try:
        # 1. 从文件名提取线索
        file_hints = extract_info_from_filename(file.filename)

        # 2. 提取简历正文文本
        full_text = extract_text_from_file(file)

        # 3. 构建带文件名提示的 Prompt
        prompt_template = """
        你是一个专业的招聘简历解析助手。
        请结合以下两部分信息，提取候选人的关键数据：
        
        【线索信息 - 来自文件命名】：
        {file_hints}
        (注意：如果简历正文中没有明确信息，请优先参考这里的线索，尤其是姓名和求职意向)
        
        【简历正文内容】：
        {text}
        
        要求：
        1. 必须返回严格的 JSON 格式。
        2. 如果某项信息在简历中找不到，请返回空字符串 "" 或空数组 []。
        3. 技能列表 (skills) 请尽量提取技术栈、编程语言、软技能等。
        
        请严格按照以下 JSON 结构输出：
        {{
            "name": "候选人姓名",
            "phone": "联系电话",
            "email": "电子邮箱",
            "title": "求职意向或当前职位",
            "education": "最高学历及学校专业",
            "summary": "个人优势或自我评价",
            "work_experience": "详细的工作经历",
            "project_experience": "主要项目经历",
            "skills": ["技能1", "技能2"],
            "resume_language": "zh-CN" 或 "en-US"
        }}
        """
        prompt = PromptTemplate.from_template(prompt_template)
        parser = JsonOutputParser()
        chain = prompt | get_llm() | parser

        # 将文件名线索和正文一起传给大模型
        input_text = full_text[:3000]
        result = chain.invoke({"file_hints": str(file_hints), "text": input_text})

        # 4. 数据清洗
        skills_raw = result.get("skills", [])
        if isinstance(skills_raw, str):
            skills = [s.strip() for s in skills_raw.split(",") if s.strip()]
        elif isinstance(skills_raw, list):
            skills = [str(s).strip() for s in skills_raw if s]
        else:
            skills = []

        # 5. 兜底逻辑：如果大模型依然没提取到，直接用文件名提取的强特征兜底
        final_name = result.get("name") or file_hints["possible_name"] or ""
        final_title = result.get("title") or file_hints["possible_title"] or ""

        return {
            "name": final_name,
            "phone": result.get("phone") or "",
            "email": result.get("email") or "",
            "title": final_title,
            "education": result.get("education") or "",
            "summary": result.get("summary") or "",
            "work_experience": result.get("work_experience") or "",
            "project_experience": result.get("project_experience") or "",
            "skills": skills,
            "resume_language": result.get("resume_language") or "zh-CN",
            "source": file.filename,
        }

    except OutputParserException as e:
        raise HTTPException(status_code=500, detail="AI 未能正确理解文档结构。")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档解析错误: {str(e)}")


# --- 3. 创建简历 (双库写入) ---
@router.post("/create", response_model=ResumeRead)
def create_resume(
    *,
    db: Session = Depends(get_db),
    resume_in: ResumeCreate,
    current_user: User = Depends(deps.get_current_user),
):
    # 1. 准备数据库对象
    db_resume = Resume(
        user_id=current_user.id,
        name=resume_in.name,
        phone=resume_in.phone,
        email=resume_in.email,
        title=resume_in.title,
        education=resume_in.education,
        # experience=resume_in.experience,
        skills=resume_in.skills,
        summary=resume_in.summary,
        work_experience=resume_in.work_experience,
        project_experience=resume_in.project_experience,
        is_default=resume_in.is_default if resume_in.is_default is not None else 0,
    )

    # 2. 写入 MySQL
    db.add(db_resume)
    db.commit()
    db.refresh(db_resume)  # 获取 ID

    # 3. 写入 Faiss 向量库
    full_text = concatenate_resume_text(resume_in)
    if full_text.strip():
        try:
            metadata = {"resume_id": db_resume.id, "title": db_resume.title}
            add_documents_to_vectorstore([full_text], [metadata])
        except Exception as e:
            print(f"Faiss 写入失败: {e}")

    return db_resume


# --- 4. 更新简历 ---
@router.put("/{resume_id}", response_model=ResumeRead)
def update_resume(
    *,
    db: Session = Depends(get_db),
    resume_id: int,
    resume_in: ResumeUpdate,
    current_user: User = Depends(deps.get_current_user),
):
    # 1. 查找并鉴权
    db_resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )

    if not db_resume:
        raise HTTPException(status_code=404, detail="简历未找到")

    # 2. 更新字段
    update_data = resume_in.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_resume, field, value)

    db.add(db_resume)
    db.commit()
    db.refresh(db_resume)

    # 3. 更新 Faiss (追加模式)
    full_text = concatenate_resume_text(resume_in)
    if full_text.strip():
        try:
            metadata = {"resume_id": db_resume.id, "title": db_resume.title}
            add_documents_to_vectorstore([full_text], [metadata])
        except Exception as e:
            print(f"Faiss 更新失败: {e}")

    return db_resume


# --- 5. 删除简历 ---
@router.delete("/{resume_id}")
def delete_resume(
    *,
    db: Session = Depends(get_db),
    resume_id: int,
    current_user: User = Depends(deps.get_current_user),
):
    obj = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )

    if not obj:
        raise HTTPException(status_code=404, detail="简历未找到")

    db.delete(obj)
    db.commit()

    # 注意：这里也应该同步删除 Faiss 中的向量，目前暂缺

    return {"detail": "删除成功"}
