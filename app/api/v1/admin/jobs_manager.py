"""
Admin 职位管理路由
包含：增删改查、文件解析（调用 LLM）、向量库同步
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status, Query
from sqlmodel import Session, select
from typing import List, Optional
from datetime import datetime
import json
import os
import shutil
import pickle
import io

# --- 文档处理依赖 ---
import pdfplumber
from docx import Document

# --- LangChain & LLM 依赖 ---
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.exceptions import OutputParserException

# --- 项目内部导入 ---
from app.models.job_position import JobPosition 
from app.schemas import job_schema  # 1. 确保导入的是更新后的 schema
from app.models.base import UserDB
from app.models import database
from app.core.deps import get_current_active_admin
from app.core.config import settings
from app.core.vector_store import get_vectorstore
from app.core.llm import get_llm

# --- 1. 配置常量 ---
UPLOAD_DIR = "uploads/jobs"
VECTOR_DB_PATH = settings.VECTOR_DB_PATH
INDEX_FILE = os.path.join(VECTOR_DB_PATH, "faiss_index.bin")
METADATA_FILE = os.path.join(VECTOR_DB_PATH, "faiss_metadata.pkl")

# 确保目录存在
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(VECTOR_DB_PATH, exist_ok=True)

# --- 2. 初始化组件 ---
llm = get_llm()  # 初始化 LLM 实例

# 初始化 Embedding 模型
local_model_path = r"E:\llm\BAAI\bge-small-zh-v1.5"
print(f"正在加载 Embedding 模型: {local_model_path} ...")
model_kwargs = {'device': 'cpu', 'trust_remote_code': True}
encode_kwargs = {'normalize_embeddings': True}
from langchain_community.embeddings import HuggingFaceEmbeddings
embeddings = HuggingFaceEmbeddings(model_name=local_model_path, model_kwargs=model_kwargs, encode_kwargs=encode_kwargs)
print("Embedding 模型加载成功！")

# 创建路由对象
router = APIRouter(prefix="/api/v1/admin", tags=["职位管理"])

# --- 3. 核心辅助函数 ---
def delete_job_from_vectorstore(job_id: int):
    """从向量库中删除指定 job_id 的记录"""
    if not os.path.exists(INDEX_FILE) or not os.path.exists(METADATA_FILE):
        return
    # 这里简化处理：只更新元数据文件
    with open(METADATA_FILE, 'rb') as f:
        metadatas = pickle.load(f)
    new_metadatas = [meta for meta in metadatas if meta.get("job_id") != str(job_id)]
    with open(METADATA_FILE, 'wb') as f:
        pickle.dump(new_metadatas, f)

def add_texts_to_vectorstore(texts: List[str], metadatas: List[dict]):
    """将文本添加到向量库"""
    if not texts:
        return
    store = get_vectorstore()
    index = store["index"]
    
    # 生成向量
    numpy_embeddings = embeddings.embed_documents(texts)
    import numpy as np
    numpy_embeddings = np.array(numpy_embeddings).astype('float32')
    index.add(numpy_embeddings)
    
    # 处理元数据
    all_metadatas = []
    if os.path.exists(METADATA_FILE):
        with open(METADATA_FILE, 'rb') as f:
            all_metadatas = pickle.load(f)
    all_metadatas.extend(metadatas)
    
    with open(METADATA_FILE, 'wb') as f:
        pickle.dump(all_metadatas, f)
    
    import faiss
    faiss.write_index(index, INDEX_FILE)

# --- 4. 新增：智能文档解析接口 ---
@router.post("/jobs/parse", response_model=job_schema.JobParseResponse)
async def parse_job_document(
    file: UploadFile = File(...),
    current_user: UserDB = Depends(get_current_active_admin)
):
    """
    解析上传的职位文档（PDF/Word/TXT），利用 LLM 提取结构化信息。
    """
    # 1. 文件类型检查
    if not file.filename.lower().endswith(('.pdf', '.docx', '.txt')):
        raise HTTPException(status_code=400, detail="仅支持 PDF, DOCX 或 TXT 文件")

    full_text = ""
    try:
        # --- 第一步：读取并提取文本 ---
        contents = await file.read()
        if file.filename.lower().endswith('.pdf'):
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
        
        # 替换原来的 docx 处理逻辑
        elif file.filename.lower().endswith('.docx'):
            doc = Document(io.BytesIO(contents))
            full_text_parts = []
            # 1. 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text_parts.append(para.text)
            # 2. 提取表格（非常重要，很多JD是表格形式）
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text_parts.append(" | ".join(row_text)) # 用分隔符隔开单元格
            full_text = "\n".join(full_text_parts)

        if not full_text.strip():
            raise HTTPException(status_code=400, detail="无法提取文档文本，请检查文件是否损坏")

        # --- 第二步：构建 LLM Chain 进行结构化提取 ---
        # 2. 定义 Prompt 模板 (更新：增加了 location, experience, education 字段)
        prompt_template = """
        你是一个专业的招聘文档解析助手。
        请仔细阅读以下文档内容，提取出职位的关键信息。
        
        要求：
        1. 如果文档中没有明确提及薪资，返回 "面议"。
        2. 技能列表 (required_skills) 请尽量提取技术栈、编程语言、工具等。
        3. 请尝试从文档中推断或提取以下信息：
           - location (工作地点): 如 东莞、深圳、远程等。
           - experience_requirement (经验要求): 如 3-5年、1-3年、无需经验等。
           - education_requirement (学历要求): 如 本科、大专、硕士等。
        
        文档内容:
        {text}
        
        请以严格的 JSON 格式输出，不要包含任何其他文字：
        {{
            "title": "职位名称",
            "company": "公司名称",
            "salary": "薪资范围",
            "location": "工作地点",
            "experience_requirement": "经验要求",
            "education_requirement": "学历要求",
            "required_skills": ["技能1", "技能2"],
            "description": "职位描述摘要（核心职责和要求）"
        }}
        """
        prompt = PromptTemplate.from_template(prompt_template)
        parser = JsonOutputParser()

        # 组装 Chain
        chain = prompt | llm | parser

        # 执行调用 (限制输入长度)
        input_text = full_text[:2500] 
        
        try:
            result = chain.invoke({"text": input_text})
        except Exception as e:
            raw_response = llm.invoke(input_text).content
            raise HTTPException(status_code=500, detail=f"AI 解析格式错误: {str(e)}")

        # --- 第三步：数据清洗 ---
        # 处理技能字段
        skills_raw = result.get("required_skills", [])
        if isinstance(skills_raw, str):
            skills = [s.strip() for s in skills_raw.split(',') if s.strip()]
        elif isinstance(skills_raw, list):
            skills = [s for s in skills_raw if s]
        else:
            skills = []

        # 处理描述字段
        desc = result.get("description", "")[:800]

        # --- 第四步：返回响应 ---
        # 3. 确保返回对象包含新字段
        return job_schema.JobParseResponse(
            title=result.get("title") or f"解析职位-{file.filename.split('.')[0]}",
            company=result.get("company") or "未知公司",
            salary=result.get("salary") or "面议",
            # 4. 新增字段映射
            location=result.get("location") or "不限",
            experience_requirement=result.get("experience_requirement") or "不限",
            education_requirement=result.get("education_requirement") or "不限",
            required_skills=skills,
            description=desc
        )

    except OutputParserException as e:
        print(f"JSON 解析异常: {e}")
        raise HTTPException(status_code=500, detail="AI 未能正确理解文档结构，请检查文档内容是否清晰。")
    except Exception as e:
        print(f"解析失败: {e}")
        raise HTTPException(status_code=500, detail=f"文档解析错误: {str(e)}")


# --- 5. 原有 CRUD 接口 (更新：适配新字段) ---
@router.get("/jobs", response_model=List[JobPosition])
def read_jobs(
    keyword: Optional[str] = Query(None, description="搜索关键词"),
    session: Session = Depends(database.get_db)
):
    statement = select(JobPosition)
    if keyword:
        statement = statement.where(
            (JobPosition.title.contains(keyword)) | 
            (JobPosition.company.contains(keyword))
        )
    result = session.execute(statement)
    jobs = result.scalars().all()
    return jobs


@router.post("/jobs", response_model=job_schema.JobRead, status_code=status.HTTP_201_CREATED)
async def create_job(
    title: str = Form(...),
    company: str = Form(...),
    salary: str = Form(...),
    # 5. 新增：接收新字段
    location: str = Form(None),
    experience_requirement: str = Form(None),
    education_requirement: str = Form(None),
    required_skills: str = Form(...),
    description: str = Form(""),
    file: Optional[UploadFile] = File(None),
    current_user: UserDB = Depends(get_current_active_admin),
    db: Session = Depends(database.get_db)
):
    file_path = None
    full_text = description

    # 处理文件上传和解析
    if file:
        file_ext = os.path.splitext(file.filename)[1]
        safe_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{title}{file_ext}"
        file_location = os.path.join(UPLOAD_DIR, safe_filename)
        with open(file_location, "wb+") as file_object:
            shutil.copyfileobj(file.file, file_object)
        file_path = file_location

        # 提取文件文本用于向量化
        try:
            if file_ext.lower() == '.pdf':
                with pdfplumber.open(file_location) as pdf:
                    full_text = ""
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            full_text += text + "\n"
            elif file_ext.lower() in ['.docx', '.doc']:
                doc = Document(file_location)
                full_text = "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            print(f"文件解析失败: {e}")
            full_text = description

    # 解析技能列表
    try:
        skills_list = json.loads(required_skills)
    except:
        skills_list = [s.strip() for s in required_skills.split(',') if s.strip()]

    # 创建数据库对象
    db_job = JobPosition(
        title=title,
        company=company,
        salary=salary,
        # 6. 新增：映射新字段
        location=location,
        experience_requirement=experience_requirement,
        education_requirement=education_requirement,
        required_skills=skills_list,
        description=full_text,
        pdf_path=file_path,
        created_at=datetime.now(),
        updated_at=datetime.now()
    )
    db.add(db_job)
    db.commit()
    db.refresh(db_job)

    # --- 向量库写入 ---
    if full_text:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        texts = text_splitter.split_text(full_text)
        metadatas = [{"id": str(db_job.id), "source": file.filename if file else "text"} for _ in texts]
        add_texts_to_vectorstore(texts, metadatas)

    return db_job


@router.put("/jobs/{id}", response_model=job_schema.JobRead)
async def update_job(
    id: int,
    title: str = Form(...),
    company: str = Form(...),
    salary: str = Form(...),
    # 7. 新增：更新接口接收新字段
    location: str = Form(None),
    experience_requirement: str = Form(None),
    education_requirement: str = Form(None),
    required_skills: str = Form(...),
    description: str = Form(""),
    file: Optional[UploadFile] = File(None),
    current_user: UserDB = Depends(get_current_active_admin),
    db: Session = Depends(database.get_db)
):
    # 查询职位
    statement = select(JobPosition).where(JobPosition.id == id)
    db_job = db.execute(statement).scalar_one_or_none()
    if not db_job:
        raise HTTPException(status_code=404, detail="职位不存在")

    full_text = description

    # 处理文件更新
    if file:
        # 删除旧文件
        if db_job.pdf_path and os.path.exists(db_job.pdf_path):
            try:
                os.remove(db_job.pdf_path)
            except:
                pass

        # 保存新文件
        file_ext = os.path.splitext(file.filename)[1]
        safe_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{title}{file_ext}"
        file_location = os.path.join(UPLOAD_DIR, safe_filename)
        with open(file_location, "wb+") as file_object:
            shutil.copyfileobj(file.file, file_object)
        db_job.pdf_path = file_location

        # 解析新文件文本
        try:
            if file_ext.lower() == '.pdf':
                with pdfplumber.open(file_location) as pdf:
                    new_full_text = ""
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            new_full_text += text + "\n"
                    full_text = new_full_text
            elif file_ext.lower() in ['.docx', '.doc']:
                doc = Document(file_location)
                full_text = "\n".join([para.text for para in doc.paragraphs])
        except:
            full_text = description
    else:
        # 无新文件，使用传入的描述或保留旧描述
        full_text = description or db_job.description

    # 更新字段
    db_job.title = title
    db_job.company = company
    db_job.salary = salary
    # 8. 新增：更新新字段
    db_job.location = location
    db_job.experience_requirement = experience_requirement
    db_job.education_requirement = education_requirement
    db_job.description = full_text
    db_job.updated_at = datetime.now()

    try:
        db_job.required_skills = json.loads(required_skills) if required_skills.startswith('[') else [s.strip() for s in required_skills.split(',') if s.strip()]
    except:
        db_job.required_skills = [s.strip() for s in required_skills.split(',') if s.strip()]

    db.add(db_job)
    db.commit()
    db.refresh(db_job)

    # --- 重新写入向量库 ---
    # 简单策略：先删除旧的（基于职位id），再写入新的
    delete_job_from_vectorstore(id)
    if full_text:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        texts = text_splitter.split_text(full_text)
        metadatas = [{"id": str(db_job.id), "source": "update"} for _ in texts]
        add_texts_to_vectorstore(texts, metadatas)

    return db_job


@router.delete("/jobs/{id}")
def delete_job(
    job_id: int,
    current_user: UserDB = Depends(get_current_active_admin),
    db: Session = Depends(database.get_db)
):
    statement = select(JobPosition).where(JobPosition.id == job_id)
    db_job = db.execute(statement).scalar_one_or_none()
    if not db_job:
        raise HTTPException(status_code=404, detail="职位不存在")

    # 删除文件
    if db_job.pdf_path and os.path.exists(db_job.pdf_path):
        os.remove(db_job.pdf_path)

    # 删除向量库记录
    delete_job_from_vectorstore(job_id)

    # 删除数据库记录
    db.delete(db_job)
    db.commit()
    return {"message": "职位删除成功"}