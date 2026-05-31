from sqlmodel import Session, select, or_
from typing import List, Optional
import logging

# 导入自定义模块
from app.models.resume import Resume
from app.schemas.resume_schema import ResumeCreate, ResumeRead, ResumeUpdate
from app.models.database import get_db

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

import io
import pdfplumber
from docx import Document
from fastapi import APIRouter, Depends, File, UploadFile, HTTPException, Query

import os
import re

from typing import List, Optional

# --- LangChain & LLM 依赖 ---
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.exceptions import OutputParserException

# --- 项目内部导入 ---
from app.models.base import UserDB
from app.core.deps import get_current_active_admin
from app.core.llm import get_llm

router = APIRouter(prefix="/api/v1/admin", tags=["简历管理"])


# --- 核心辅助函数：从文件名中提取关键信息 ---
def extract_info_from_filename(filename: str) -> dict:
    """
    从文件名中提取姓名、职位、薪资等信息
    支持常见的命名格式，如：张三_前端开发_15k.pdf、李四-Java工程师-3年经验.docx
    """
    # 去掉文件后缀
    name_without_ext = os.path.splitext(filename)[0]
    # 替换常见的分隔符（下划线、横杠、空格）为统一的空格
    clean_name = re.sub(r"[_\-\s]+", " ", name_without_ext).strip()

    extracted_info = {
        "filename_hint": clean_name,
        "possible_name": "",
        "possible_title": "",
        "possible_salary": "",
    }

    # 1. 提取薪资（匹配如：15k, 20K, 15-25k, 20k-30k 等）
    salary_match = re.search(
        r"(\d+\.?\d*)\s*[-~至到]?\s*(\d*)\s*k", clean_name, re.IGNORECASE
    )
    if salary_match:
        min_s = salary_match.group(1)
        max_s = salary_match.group(2)
        if max_s:
            extracted_info["possible_salary"] = f"{min_s}k-{max_s}k"
        else:
            extracted_info["possible_salary"] = f"{min_s}k"

    # 2. 提取职位（这里做一个简单的启发式匹配，实际可以根据你的业务扩展关键词库）
    # 匹配文件名中常见的职位关键词
    job_keywords = [
        "前端",
        "后端",
        "Java",
        "Python",
        "算法",
        "产品",
        "运营",
        "测试",
        "UI",
        "设计",
        "工程师",
        "开发",
        "经理",
        "总监",
        "实习",
    ]
    found_jobs = [kw for kw in job_keywords if kw.lower() in clean_name.lower()]
    if found_jobs:
        # 简单取第一个匹配到的，或者你可以写更复杂的逻辑提取完整职位短语
        extracted_info["possible_title"] = found_jobs[0]

    # 3. 提取姓名（假设文件名开头2-4个字符是中文姓名）
    name_match = re.match(r"^([\u4e00-\u9fa5]{2,4})", clean_name)
    if name_match:
        extracted_info["possible_name"] = name_match.group(1)

    return extracted_info


# --- 辅助函数：提取文档纯文本（复用之前的逻辑） ---
def extract_text_from_file(file: UploadFile) -> str:
    contents = file.file.read()
    full_text = ""
    try:
        if file.filename.lower().endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(contents)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"
        elif file.filename.lower().endswith(".docx"):
            doc = Document(io.BytesIO(contents))
            full_text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    ]
                    if row_text:
                        full_text_parts.append(" | ".join(row_text))
            full_text = "\n".join(full_text_parts)
        elif file.filename.lower().endswith(".txt"):
            full_text = contents.decode("utf-8")
        if not full_text.strip():
            raise ValueError("无法提取文档文本")
        return full_text
    finally:
        file.file.seek(0)


# --- 简历智能解析接口（升级版） ---
@router.post("/resumes/parse")
async def parse_resume_document(
    file: UploadFile = File(...),
    current_user: UserDB = Depends(get_current_active_admin),
):
    if not file.filename.lower().endswith((".pdf", ".docx", ".txt")):
        raise HTTPException(status_code=400, detail="仅支持 PDF, DOCX 或 TXT 文件")

    try:
        # 1. 从文件名提取线索
        file_hints = extract_info_from_filename(file.filename)

        # 2. 提取简历正文文本
        full_text = extract_text_from_file(file)

        # 3. 构建带文件名提示的 Prompt
        prompt_template = """
        你是一个专业的招聘简历解析助手。
        请结合以下两部分信息，提取候选人的关键数据：
        
        【线索信息 - 来自文件命名】：
        {file_hints}
        (注意：如果简历正文中没有明确信息，请优先参考这里的线索，尤其是姓名和求职意向)
        
        【简历正文内容】：
        {text}}
        
        要求：
        1. 必须返回严格的 JSON 格式。
        2. 如果某项信息在简历中找不到，请返回空字符串 "" 或空数组 []。
        3. 技能列表 (skills) 请尽量提取技术栈、编程语言、软技能等。
        
        请严格按照以下 JSON 结构输出：
        {{
            "name": "候选人姓名",
            "phone": "联系电话",
            "email": "电子邮箱",
            "title": "求职意向或当前职位",
            "education": "最高学历及学校专业",
            "summary": "个人优势或自我评价",
            "work_experience": "详细的工作经历",
            "project_experience": "主要项目经历",
            "skills": ["技能1", "技能2"],
            "resume_language": "zh-CN" 或 "en-US"
        }}
        """
        prompt = PromptTemplate.from_template(prompt_template)
        parser = JsonOutputParser()
        chain = prompt | get_llm() | parser

        # 将文件名线索和正文一起传给大模型
        input_text = full_text[:3000]
        result = chain.invoke({"file_hints": str(file_hints), "text": input_text})

        # 4. 数据清洗
        skills_raw = result.get("skills", [])
        if isinstance(skills_raw, str):
            skills = [s.strip() for s in skills_raw.split(",") if s.strip()]
        elif isinstance(skills_raw, list):
            skills = [str(s).strip() for s in skills_raw if s]
        else:
            skills = []

        # 5. 兜底逻辑：如果大模型依然没提取到，直接用文件名提取的强特征兜底
        final_name = result.get("name") or file_hints["possible_name"] or ""
        final_title = result.get("title") or file_hints["possible_title"] or ""

        return {
            "name": final_name,
            "phone": result.get("phone") or "",
            "email": result.get("email") or "",
            "title": final_title,
            "education": result.get("education") or "",
            "summary": result.get("summary") or "",
            "work_experience": result.get("work_experience") or "",
            "project_experience": result.get("project_experience") or "",
            "skills": skills,
            "resume_language": result.get("resume_language") or "zh-CN",
            "source": file.filename,
        }

    except OutputParserException as e:
        raise HTTPException(status_code=500, detail="AI 未能正确理解文档结构。")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文档解析错误: {str(e)}")


# 1. 获取简历列表 + 智能搜索 + 分页
@router.get("/resumes", response_model=List[ResumeRead])
def list_resumes(
    *,
    session: Session = Depends(get_db),
    # 搜索参数
    q: Optional[str] = Query(None, description="搜索关键词（姓名、邮箱或技能）"),
    # 分页参数
    offset: int = Query(0, ge=0, description="分页偏移量"),
    limit: int = Query(10, ge=1, le=100, description="每页数量"),
):
    """
    获取简历列表，支持模糊搜索和分页
    """
    try:
        # 构建基础查询语句
        statement = select(Resume)

        # 如果有关键词，添加模糊查询条件
        if q:
            # 构建模糊匹配逻辑：匹配 name, email 或 skills (JSON字段)
            search_filter = or_(
                Resume.name.contains(q),
                Resume.email.contains(q),
                Resume.skills.contains(q),  # 注意：这里依赖数据库的JSON文本匹配能力
            )
            statement = statement.where(search_filter)

        # 排序：按创建时间倒序
        statement = statement.order_by(Resume.created_at.desc())

        # 分页处理
        statement = statement.offset(offset).limit(limit)

        # 执行查询
        result = session.execute(statement)
        resumes = result.scalars().all()

        # 获取总数（用于前端分页）
        count_statement = select(Resume)
        if q:
            count_statement = count_statement.where(search_filter)
        count_result = session.execute(count_statement)
        total = len(count_result.scalars().all())

        logger.info(f"查询简历列表成功，偏移量={offset}, 限制={limit}, 总数={total}")
        return resumes

    except Exception as e:
        logger.error(f"获取简历列表失败: {e}")
        raise HTTPException(status_code=500, detail="服务器内部错误")


# 2. 新增简历
@router.post("/resumes", response_model=ResumeRead)
def create_resume(*, session: Session = Depends(get_db), resume_in: ResumeCreate):
    try:
        # 转换模型
        # 使用 model_dump() 替代旧的 dict()
        resume_data = resume_in.model_dump(exclude_unset=True)

        # 确保状态符合数据库默认值 (active)
        if not resume_data.get("status"):
            resume_data["status"] = "active"

        db_resume = Resume(**resume_data)

        session.add(db_resume)
        session.commit()
        session.refresh(db_resume)

        logger.info(f"新增简历成功: ID={db_resume.id}, 姓名={db_resume.name}")
        return db_resume

    except Exception as e:
        session.rollback()
        logger.error(f"新增简历失败: {e}")
        raise HTTPException(status_code=500, detail="新增简历失败")


# 3. 获取简历详情
@router.get("/resumes/{resume_id}", response_model=ResumeRead)
def get_resume_detail(*, session: Session = Depends(get_db), resume_id: int):
    resume = session.get(Resume, resume_id)
    if not resume:
        raise HTTPException(status_code=404, detail="简历未找到")
    return resume


# 4. 更新简历
@router.put("/resumes/{resume_id}", response_model=ResumeRead)
def update_resume(
    *, session: Session = Depends(get_db), resume_id: int, resume_in: ResumeUpdate
):
    db_resume = session.get(Resume, resume_id)
    if not db_resume:
        raise HTTPException(status_code=404, detail="简历未找到")

    try:
        # 获取传入的数据，排除未设置的字段
        update_data = resume_in.model_dump(exclude_unset=True)

        # 遍历字典更新对象属性
        for key, value in update_data.items():
            setattr(db_resume, key, value)

        session.add(db_resume)
        session.commit()
        session.refresh(db_resume)

        logger.info(f"更新简历成功: ID={db_resume.id}")
        return db_resume

    except Exception as e:
        session.rollback()
        logger.error(f"更新简历失败: {e}")
        raise HTTPException(status_code=500, detail="更新失败")


# 5. 删除简历
@router.delete("/resumes/{resume_id}")
def delete_resume(*, session: Session = Depends(get_db), resume_id: int):
    resume = session.get(Resume, resume_id)
    if not resume:
        raise HTTPException(status_code=404, detail="简历未找到")

    try:
        session.delete(resume)
        session.commit()
        logger.info(f"删除简历成功: ID={resume_id}")
        return {"message": "删除成功"}

    except Exception as e:
        session.rollback()
        logger.error(f"删除简历失败: {e}")
        raise HTTPException(status_code=500, detail="删除失败")
